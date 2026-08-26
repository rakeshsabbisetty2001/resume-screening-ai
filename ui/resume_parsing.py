"""Extracts plain text from an uploaded resume file (.txt/.pdf/.docx).

Kept dependency-free of streamlit itself so this can be unit-checked on its
own — see the __main__ self-check below. Only .docx is supported for Word
files, not the legacy binary .doc format: python-docx can't read .doc at
all (it's a pre-XML binary format), and pulling in something that can
(e.g. antiword, a system binary) is a lot of deploy-fragile weight for a
format that's been obsolete since Word 2007. Declared honestly in the UI's
file-type filter rather than silently accepting .doc and failing on it.
"""
import io

from docx import Document
from pypdf import PdfReader


class ExtractionError(Exception):
    """Raised when a file can't be parsed, or parses to no usable text
    (e.g. a scanned/image-only PDF with no text layer)."""


def extract_text(filename: str, data: bytes) -> str:
    ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""

    if ext == "txt":
        text = data.decode("utf-8", errors="replace")
    elif ext == "pdf":
        try:
            reader = PdfReader(io.BytesIO(data))
            text = "\n".join(page.extract_text() or "" for page in reader.pages)
        except Exception as e:
            raise ExtractionError(f"couldn't read this PDF ({e})") from None
    elif ext == "docx":
        try:
            doc = Document(io.BytesIO(data))
            text = "\n".join(p.text for p in doc.paragraphs)
        except Exception as e:
            raise ExtractionError(f"couldn't read this .docx ({e})") from None
    else:
        raise ExtractionError(f"unsupported file type: .{ext}")

    if not text.strip():
        raise ExtractionError(
            "no text could be extracted — likely a scanned/image-only file"
        )
    return text


def _build_minimal_pdf(content_stream: bytes) -> bytes:
    """Assembles a structurally valid one-page PDF with a real xref table,
    for the self-check below — self-contained so the check has no
    dependency on a fixture file or reportlab."""
    objects = [
        b"<</Type/Catalog/Pages 2 0 R>>",
        b"<</Type/Pages/Kids[3 0 R]/Count 1>>",
        b"<</Type/Page/Parent 2 0 R/MediaBox[0 0 200 200]"
        b"/Resources<</Font<</F1 4 0 R>>>>/Contents 5 0 R>>",
        b"<</Type/Font/Subtype/Type1/BaseFont/Helvetica>>",
        b"<</Length " + str(len(content_stream)).encode() + b">>stream\n"
        + content_stream + b"\nendstream",
    ]
    out = bytearray(b"%PDF-1.4\n")
    offsets = []
    for i, body in enumerate(objects, start=1):
        offsets.append(len(out))
        out += f"{i} 0 obj".encode() + body + b"endobj\n"
    xref_start = len(out)
    out += f"xref\n0 {len(objects) + 1}\n0000000000 65535 f \n".encode()
    for off in offsets:
        out += f"{off:010d} 00000 n \n".encode()
    out += (
        f"trailer<</Size {len(objects) + 1}/Root 1 0 R>>\n".encode()
        + b"startxref\n" + str(xref_start).encode() + b"\n%%EOF"
    )
    return bytes(out)


def demo():
    # .txt
    assert extract_text("r.txt", b"hello world").strip() == "hello world"

    # .docx — build a real one in memory, round-trip it
    buf = io.BytesIO()
    doc = Document()
    doc.add_paragraph("Jane Doe")
    doc.add_paragraph("Software Engineer, 5 years Python")
    doc.save(buf)
    text = extract_text("r.docx", buf.getvalue())
    assert "Jane Doe" in text and "Python" in text

    # .pdf — build a real, structurally valid minimal PDF (pypdf needs a
    # real xref table, unlike some lenient readers) with byte offsets
    # computed rather than hand-counted, so a stray edit above can't
    # silently desync them.
    pdf_bytes = _build_minimal_pdf(b"BT /F1 12 Tf 10 100 Td (John Smith Resume) Tj ET")
    text = extract_text("r.pdf", pdf_bytes)
    assert "John Smith" in text

    # empty/unreadable -> ExtractionError, not a silent empty string
    try:
        extract_text("r.pdf", b"not a real pdf")
        raise AssertionError("expected ExtractionError")
    except ExtractionError:
        pass

    # unsupported extension
    try:
        extract_text("r.doc", b"anything")
        raise AssertionError("expected ExtractionError")
    except ExtractionError:
        pass

    print("resume_parsing self-check: OK")


if __name__ == "__main__":
    demo()
